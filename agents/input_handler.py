from docx import Document
import fitz
from schemas.graph_state import State
import io


def read_file(state: State):
    response = {"candidate_name": state["candidate_name"]}
    for key, file_path in state["files"].items():
        if type(file_path) == str:
            if file_path.endswith(".pdf"):
                doc = fitz.open(file_path)
                response.update({key: "\n".join([page.get_text() for page in doc])})
            elif file_path.endswith(".docx"):
                response.update(
                    {key: "\n".join([p.text for p in Document(file_path).paragraphs])}
                )
            elif key == "candidate_resume":
                raise ValueError("Unsupported file format")
            else:
                response.update({key: file_path})
        else:
            if file_path["filename"].endswith(".pdf"):
                contents = file_path["content"]
                doc = fitz.open(
                    stream=contents,
                    filetype="pdf",
                )
                response.update({key: "\n".join([page.get_text() for page in doc])})
            elif file_path["filename"].endswith(".docx"):
                # doc = file_path["content"].file.read()
                response.update(
                    {
                        key: "\n".join(
                            [
                                p.text
                                for p in Document(
                                    io.BytesIO(file_path["content"])
                                ).paragraphs
                            ]
                        )
                    }
                )
            elif key == "candidate_resume":
                raise ValueError("Unsupported file format")
            else:
                response.update({key: file_path})

    return response
